#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Versão avançada do script para geração de relatórios de manutenção.
Permite personalização completa dos dados via parâmetros, mantendo
a formatação idêntica ao modelo report.doc original.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime
import argparse
import json

def definir_estilo_paragrafo(paragrafo, espacamento_antes=0, espacamento_depois=0, 
                            espacamento_linha=1.0, alinhamento=WD_PARAGRAPH_ALIGNMENT.LEFT):
    """Define o estilo completo do parágrafo."""
    paragrafo.alignment = alinhamento
    paragrafo_formato = paragrafo.paragraph_format
    paragrafo_formato.space_before = Pt(espacamento_antes)
    paragrafo_formato.space_after = Pt(espacamento_depois)
    paragrafo_formato.line_spacing = espacamento_linha

def definir_estilo_fonte(run, tamanho=11, negrito=False, italico=False, nome_fonte="Arial"):
    """Define o estilo da fonte."""
    run.font.name = nome_fonte
    run.font.size = Pt(tamanho)
    run.font.bold = negrito
    run.font.italic = italico

def adicionar_borda_tabela(tabela, largura_borda=1):
    """Adiciona borda à tabela com espessura específica."""
    tbl = tabela._tbl
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), str(largura_borda * 4))  # 4 = 1pt
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)
        tcPr.append(tcBorders)

def configurar_larguras_celulas(tabela, larguras):
    """Define larguras específicas para cada coluna da tabela."""
    for i, largura in enumerate(larguras):
        for celula in tabela.columns[i].cells:
            celula.width = Inches(largura)

def gerar_relatorio(dados):
    """
    Gera um relatório de manutenção personalizado no formato Word,
    com formatação idêntica ao modelo original.
    
    Args:
        dados: Dicionário com os dados do relatório
    
    Returns:
        str: Caminho do arquivo gerado
    """
    # Verificar se os dados necessários estão presentes
    if not isinstance(dados, dict):
        raise ValueError("Os dados devem ser fornecidos como um dicionário")
        
    # Valores padrão para dados ausentes
    dados_padrao = {
        "titulo": "RELATÓRIO DE MANUTENÇÃO",
        "data": datetime.now().strftime("%d/%m/%Y"),
        "equipamento": {
            "nome": "Não especificado",
            "numero_serie": "Não especificado",
            "localizacao": "Não especificado"
        },
        "descricao_servico": "Não especificado",
        "pecas": [],
        "observacoes": "",
        "tecnico": {
            "nome": "Técnico Responsável",
            "cargo": "Técnico de Manutenção"
        },
        "nome_arquivo": "relatorio_personalizado.docx"
    }
    
    # Mesclar dados padrão com dados fornecidos
    for chave, valor in dados_padrao.items():
        if chave not in dados:
            dados[chave] = valor
        elif isinstance(valor, dict) and isinstance(dados[chave], dict):
            for sub_chave, sub_valor in valor.items():
                if sub_chave not in dados[chave]:
                    dados[chave][sub_chave] = sub_valor
    
    # Criar documento
    doc = Document()
    
    # Configurar margens da página (em centímetros)
    secao = doc.sections[0]
    secao.left_margin = Cm(2.5)
    secao.right_margin = Cm(2.5)
    secao.top_margin = Cm(2.5)
    secao.bottom_margin = Cm(2.0)
    
    # Propriedades do documento
    propriedades = doc.core_properties
    propriedades.author = "Sistema de Relatórios de Manutenção"
    propriedades.title = "Relatório de Manutenção"
    
    # TÍTULO DO RELATÓRIO
    titulo = doc.add_heading(dados["titulo"], level=0)
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    titulo.paragraph_format.space_after = Pt(18)
    for run in titulo.runs:
        definir_estilo_fonte(run, tamanho=16, negrito=True)
    
    # DATA DO RELATÓRIO - alinhada à direita com formatação específica
    data_paragrafo = doc.add_paragraph()
    data_paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    data_paragrafo.paragraph_format.space_before = Pt(0)
    data_paragrafo.paragraph_format.space_after = Pt(24)
    data_texto = data_paragrafo.add_run(f'Data: {dados["data"]}')
    definir_estilo_fonte(data_texto, tamanho=11, negrito=True)
    
    # SEÇÃO: INFORMAÇÕES DO EQUIPAMENTO
    # Título da seção
    info_titulo = doc.add_heading('Informações do Equipamento', level=1)
    definir_estilo_paragrafo(info_titulo, espacamento_depois=12)
    for run in info_titulo.runs:
        definir_estilo_fonte(run, tamanho=14, negrito=True)
    
    # Tabela de informações com formatação precisa
    tabela_info = doc.add_table(rows=3, cols=2)
    tabela_info.style = 'Table Grid'
    tabela_info.autofit = False
    configurar_larguras_celulas(tabela_info, [2.0, 4.0])
    
    # Preenchimento e formatação das células
    pares_info = [
        ("Equipamento:", dados["equipamento"]["nome"]),
        ("Número de Série:", dados["equipamento"]["numero_serie"]),
        ("Localização:", dados["equipamento"]["localizacao"])
    ]
    
    for i, (chave, valor) in enumerate(pares_info):
        # Formatação da célula de chave
        celula_chave = tabela_info.cell(i, 0)
        celula_chave.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        celula_chave.paragraphs[0].paragraph_format.space_before = Pt(6)
        celula_chave.paragraphs[0].paragraph_format.space_after = Pt(6)
        celula_chave.paragraphs[0].paragraph_format.left_indent = Pt(6)
        run_chave = celula_chave.paragraphs[0].add_run(chave)
        definir_estilo_fonte(run_chave, tamanho=11, negrito=True)
        
        # Formatação da célula de valor
        celula_valor = tabela_info.cell(i, 1)
        celula_valor.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        celula_valor.paragraphs[0].paragraph_format.space_before = Pt(6)
        celula_valor.paragraphs[0].paragraph_format.space_after = Pt(6)
        celula_valor.paragraphs[0].paragraph_format.left_indent = Pt(6)
        run_valor = celula_valor.paragraphs[0].add_run(valor)
        definir_estilo_fonte(run_valor, tamanho=11)
    
    # Adicionar borda à tabela com a espessura exata
    adicionar_borda_tabela(tabela_info, largura_borda=1)
    
    # Espaçamento após a tabela
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # SEÇÃO: DESCRIÇÃO DO SERVIÇO
    # Título da seção
    desc_titulo = doc.add_heading('Descrição do Serviço', level=1)
    definir_estilo_paragrafo(desc_titulo, espacamento_depois=12)
    for run in desc_titulo.runs:
        definir_estilo_fonte(run, tamanho=14, negrito=True)
    
    # Texto da descrição com formatação específica
    desc_paragrafo = doc.add_paragraph()
    definir_estilo_paragrafo(desc_paragrafo, espacamento_depois=12)
    desc_texto = desc_paragrafo.add_run(dados["descricao_servico"])
    definir_estilo_fonte(desc_texto, tamanho=11)
    
    # SEÇÃO: PEÇAS UTILIZADAS
    if dados["pecas"]:
        # Título da seção
        pecas_titulo = doc.add_heading('Peças Utilizadas', level=1)
        definir_estilo_paragrafo(pecas_titulo, espacamento_depois=12)
        for run in pecas_titulo.runs:
            definir_estilo_fonte(run, tamanho=14, negrito=True)
        
        # Tabela de peças com formatação específica
        tabela_pecas = doc.add_table(rows=1, cols=3)
        tabela_pecas.style = 'Table Grid'
        tabela_pecas.autofit = False
        configurar_larguras_celulas(tabela_pecas, [3.0, 1.0, 2.0])
        
        # Cabeçalho da tabela
        cabecalhos = ["Nome", "Quantidade", "Observações"]
        for i, texto in enumerate(cabecalhos):
            celula = tabela_pecas.cell(0, i)
            celula.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragrafo = celula.paragraphs[0]
            paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragrafo.paragraph_format.space_before = Pt(6)
            paragrafo.paragraph_format.space_after = Pt(6)
            run = paragrafo.add_run(texto)
            definir_estilo_fonte(run, tamanho=11, negrito=True)
        
        # Adicionar linhas à tabela com formatação precisa
        for peca in dados["pecas"]:
            # Adicionar nova linha
            linha = tabela_pecas.add_row()
            
            # Formatar célula Nome
            celula_nome = linha.cells[0]
            celula_nome.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragrafo_nome = celula_nome.paragraphs[0]
            paragrafo_nome.paragraph_format.space_before = Pt(6)
            paragrafo_nome.paragraph_format.space_after = Pt(6)
            paragrafo_nome.paragraph_format.left_indent = Pt(6)
            run_nome = paragrafo_nome.add_run(peca.get("nome", ""))
            definir_estilo_fonte(run_nome, tamanho=11)
            
            # Formatar célula Quantidade
            celula_qtd = linha.cells[1]
            celula_qtd.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragrafo_qtd = celula_qtd.paragraphs[0]
            paragrafo_qtd.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragrafo_qtd.paragraph_format.space_before = Pt(6)
            paragrafo_qtd.paragraph_format.space_after = Pt(6)
            run_qtd = paragrafo_qtd.add_run(str(peca.get("quantidade", "")))
            definir_estilo_fonte(run_qtd, tamanho=11)
            
            # Formatar célula Observações
            celula_obs = linha.cells[2]
            celula_obs.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragrafo_obs = celula_obs.paragraphs[0]
            paragrafo_obs.paragraph_format.space_before = Pt(6)
            paragrafo_obs.paragraph_format.space_after = Pt(6)
            paragrafo_obs.paragraph_format.left_indent = Pt(6)
            run_obs = paragrafo_obs.add_run(peca.get("observacoes", ""))
            definir_estilo_fonte(run_obs, tamanho=11)
        
        # Adicionar borda à tabela
        adicionar_borda_tabela(tabela_pecas, largura_borda=1)
        
        # Espaçamento após a tabela
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # SEÇÃO: OBSERVAÇÕES ADICIONAIS
    if dados["observacoes"]:
        # Título da seção
        obs_titulo = doc.add_heading('Observações Adicionais', level=1)
        definir_estilo_paragrafo(obs_titulo, espacamento_depois=12)
        for run in obs_titulo.runs:
            definir_estilo_fonte(run, tamanho=14, negrito=True)
        
        # Texto das observações
        obs_paragrafo = doc.add_paragraph()
        definir_estilo_paragrafo(obs_paragrafo, espacamento_depois=12)
        obs_texto = obs_paragrafo.add_run(dados["observacoes"])
        definir_estilo_fonte(obs_texto, tamanho=11)
    
    # SEÇÃO: ASSINATURA
    # Espaçamento antes da assinatura
    doc.add_paragraph().paragraph_format.space_after = Pt(36)
    
    # Linha de assinatura
    assinatura_linha = doc.add_paragraph()
    assinatura_linha.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    assinatura_linha.paragraph_format.space_after = Pt(0)
    assinatura_linha.add_run("____________________________").font.size = Pt(11)
    
    # Nome do técnico
    nome_tecnico = doc.add_paragraph()
    nome_tecnico.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    nome_tecnico.paragraph_format.space_before = Pt(6)
    nome_tecnico.paragraph_format.space_after = Pt(0)
    run_nome = nome_tecnico.add_run(dados["tecnico"]["nome"])
    definir_estilo_fonte(run_nome, tamanho=11, negrito=True)
    
    # Cargo do técnico
    cargo_tecnico = doc.add_paragraph()
    cargo_tecnico.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cargo_tecnico.paragraph_format.space_before = Pt(0)
    run_cargo = cargo_tecnico.add_run(dados["tecnico"]["cargo"])
    definir_estilo_fonte(run_cargo, tamanho=11)
    
    # Salvar o documento
    caminho_arquivo = dados["nome_arquivo"]
    doc.save(caminho_arquivo)
    caminho_absoluto = os.path.abspath(caminho_arquivo)
    print(f"Relatório gerado com sucesso: {caminho_absoluto}")
    
    return caminho_absoluto

def carregar_json(caminho_arquivo):
    """Carrega dados de um arquivo JSON."""
    with open(caminho_arquivo, 'r', encoding='utf-8') as f:
        return json.load(f)

def main():
    """Função principal para execução via linha de comando."""
    parser = argparse.ArgumentParser(description='Gera relatórios de manutenção personalizados.')
    parser.add_argument('--json', help='Caminho para arquivo JSON com dados do relatório')
    parser.add_argument('--output', help='Nome do arquivo de saída (opcional)')
    args = parser.parse_args()
    
    # Dados de exemplo
    exemplo_dados = {
        "titulo": "RELATÓRIO DE MANUTENÇÃO",
        "data": datetime.now().strftime("%d/%m/%Y"),
        "equipamento": {
            "nome": "Servidor Dell PowerEdge R740",
            "numero_serie": "SRV78945612",
            "localizacao": "Sala de Servidores - Andar 2"
        },
        "descricao_servico": (
            "Foi realizada uma manutenção preventiva no servidor, incluindo limpeza dos "
            "componentes internos, verificação do sistema de refrigeração e atualização "
            "do firmware. Também foram substituídos dois discos rígidos que apresentavam "
            "erros nos testes de diagnóstico S.M.A.R.T."
        ),
        "pecas": [
            {"nome": "Disco Rígido SAS 600GB", "quantidade": 2, "observacoes": "Substituição preventiva"},
            {"nome": "Pasta Térmica", "quantidade": 1, "observacoes": "Aplicada no processador"},
            {"nome": "Filtro de Ar", "quantidade": 1, "observacoes": "Substituído por estar saturado"}
        ],
        "observacoes": (
            "Recomenda-se agendar uma manutenção completa do sistema de refrigeração "
            "nos próximos 3 meses, pois foram observados sinais de desgaste nas "
            "ventoinhas principais. A temperatura de operação está dentro dos parâmetros "
            "normais, mas é importante realizar esse serviço preventivamente."
        ),
        "tecnico": {
            "nome": "Carlos Eduardo Silva",
            "cargo": "Técnico de Manutenção"
        },
        "nome_arquivo": "relatorio_teste_avancado.docx"
    }
    
    # Carregar dados do arquivo JSON se especificado
    if args.json:
        dados = carregar_json(args.json)
    else:
        dados = exemplo_dados
    
    # Sobrescrever nome do arquivo se especificado
    if args.output:
        dados["nome_arquivo"] = args.output
    
    # Gerar relatório
    gerar_relatorio(dados)

if __name__ == "__main__":
    main() 