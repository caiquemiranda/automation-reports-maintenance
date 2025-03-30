import os
import shutil
import re
import sys
import subprocess
import tempfile
from pathlib import Path
from datetime import datetime

# Verificar se pywin32 está instalado
try:
    import win32com.client
    import win32com.client.dynamic
    from win32com.client import constants
    HAS_PYWIN32 = True
except ImportError:
    HAS_PYWIN32 = False
    print("Aviso: A biblioteca 'pywin32' não está instalada. Funcionalidades de edição serão limitadas.")
    print("Para instalar, execute: pip install pywin32")

# Definir os marcadores que serão substituídos
MARCADORES = [
    "[Data da Solicitação]",
    "[Data de Execução]",
    "[Localização]",
    "[O.S]",
    "[SERVIÇO A REALIZAR]",
    "[OBSERVAÇÃO]"
]

# Lista de campos que devem estar em negrito
CAMPOS_NEGRITO = [
    "Código da Manutenção:",
    "Data da Solicitação:",
    "Código do equipamento:",
    "Data de Execução:",
    "Nome do Equipamento:",
    "Prioridade:",
    "Localização:",
    "Requisitante:",
    "Número da O.S.:",
    "Centro de custo:",
    "SERVIÇO:",
    "OBSERVAÇÃO:"
]

# Campos que devem ser alinhados à direita (coluna direita)
CAMPOS_DIREITA = [
    "Data da Solicitação:",
    "Data de Execução:",
    "Prioridade:",
    "Requisitante:",
    "Centro de custo:"
]

def replicar_documento(caminho_origem, caminho_destino):
    """
    Replica o documento de origem para o de destino usando cópia binária
    
    Args:
        caminho_origem: Caminho do arquivo de origem
        caminho_destino: Caminho do arquivo de destino
        
    Returns:
        Boolean indicando se a operação foi bem-sucedida
    """
    try:
        # Realizar cópia direta (binária) para garantir precisão
        print(f"Realizando cópia binária do arquivo {caminho_origem}...")
        shutil.copy2(caminho_origem, caminho_destino)
        print(f"Arquivo copiado com sucesso para: {caminho_destino}")
        return True
    except Exception as e:
        print(f"Erro ao replicar o documento: {e}")
        return False

def verificar_conteudo(caminho_origem, caminho_destino):
    """
    Verifica se o conteúdo dos arquivos é idêntico
    
    Args:
        caminho_origem: Caminho do arquivo de origem
        caminho_destino: Caminho do arquivo de destino
        
    Returns:
        Boolean indicando se os arquivos são idênticos
    """
    try:
        with open(caminho_origem, 'rb') as f1, open(caminho_destino, 'rb') as f2:
            conteudo1 = f1.read()
            conteudo2 = f2.read()
            
            if conteudo1 == conteudo2:
                print("Verificação: Os arquivos são IDÊNTICOS!")
                return True
            else:
                print("Verificação: Os arquivos são DIFERENTES!")
                tamanho1 = len(conteudo1)
                tamanho2 = len(conteudo2)
                print(f"Tamanho do arquivo original: {tamanho1} bytes")
                print(f"Tamanho do arquivo modificado: {tamanho2} bytes")
                
                if tamanho1 == tamanho2:
                    print("Os arquivos têm o mesmo tamanho, mas conteúdo diferente")
                else:
                    print(f"Diferença de tamanho: {abs(tamanho1 - tamanho2)} bytes")
                
                return False
    except Exception as e:
        print(f"Erro ao verificar conteúdo: {e}")
        return False

def aplicar_formatacao_word(caminho_arquivo):
    """
    Aplica formatação (negrito e alinhamento) aos campos específicos usando Word
    
    Args:
        caminho_arquivo: Caminho do arquivo a ser formatado
        
    Returns:
        Boolean indicando se a formatação foi bem-sucedida
    """
    if not HAS_PYWIN32:
        print("Erro: A biblioteca 'pywin32' é necessária para esta função.")
        return False
    
    caminho_completo = os.path.abspath(caminho_arquivo)
    
    try:
        # Inicializar Word
        print(f"Iniciando Microsoft Word para formatar {caminho_arquivo}...")
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False  # Execução invisível
        
        # Abrir o documento
        doc = word.Documents.Open(caminho_completo)
        
        # Percorrer o documento procurando os campos para aplicar negrito e alinhar
        for campo in CAMPOS_NEGRITO:
            # Configurar a busca
            word.Selection.HomeKey(Unit=constants.wdStory)  # Ir para o início do documento
            word.Selection.Find.ClearFormatting()
            word.Selection.Find.Text = campo
            word.Selection.Find.Forward = True
            word.Selection.Find.MatchCase = True
            word.Selection.Find.MatchWholeWord = False
            
            # Procurar todas as ocorrências e aplicar negrito
            while word.Selection.Find.Execute():
                word.Selection.Font.Bold = True
                
                # Aplicar alinhamento de acordo com o campo
                if campo in CAMPOS_DIREITA:
                    # Para campos da coluna direita, alinhar o parágrafo à direita
                    try:
                        # Selecionar o parágrafo inteiro que contém o campo
                        word.Selection.EndKey(Unit=constants.wdLine)
                        word.Selection.HomeKey(Unit=constants.wdLine, Extend=True)
                        word.Selection.ParagraphFormat.Alignment = constants.wdAlignParagraphRight
                    except:
                        print(f"Não foi possível alinhar à direita: {campo}")
                else:
                    # Para campos da coluna esquerda, alinhar o parágrafo à esquerda
                    try:
                        word.Selection.ParagraphFormat.Alignment = constants.wdAlignParagraphLeft
                    except:
                        print(f"Não foi possível alinhar à esquerda: {campo}")
        
        # Salvar o documento
        doc.Save()
        
        # Fechar o documento
        doc.Close()
        
        # Encerrar o Word
        word.Quit()
        
        print("Formatação aplicada com sucesso!")
        return True
    
    except Exception as e:
        print(f"Erro ao aplicar formatação com Word: {e}")
        
        # Em caso de erro, garantir que o Word seja fechado
        try:
            if 'doc' in locals() and doc:
                doc.Close(SaveChanges=False)
            if 'word' in locals() and word:
                word.Quit()
        except:
            pass
        
        return False

def modificar_doc_com_word(caminho_arquivo, substituicoes):
    """
    Modifica um documento .doc usando a API do Word (via COM)
    
    Args:
        caminho_arquivo: Caminho do arquivo a ser modificado
        substituicoes: Dicionário com pares de marcadores e seus novos valores
        
    Returns:
        Boolean indicando se a modificação foi bem-sucedida
    """
    if not HAS_PYWIN32:
        print("Erro: A biblioteca 'pywin32' é necessária para esta função.")
        return False
    
    caminho_completo = os.path.abspath(caminho_arquivo)
    
    try:
        # Inicializar Word
        print(f"Iniciando Microsoft Word para editar {caminho_arquivo}...")
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False  # Execução invisível
        
        # Abrir o documento
        doc = word.Documents.Open(caminho_completo)
        
        # Realizar as substituições para cada marcador
        total_substituicoes = 0
        
        for marcador, novo_valor in substituicoes.items():
            # Limpar formatação de busca e substituição
            word.Selection.Find.ClearFormatting()
            word.Selection.Find.Replacement.ClearFormatting()
            
            # Configurar a busca
            find_obj = word.Selection.Find
            find_obj.Text = marcador
            find_obj.Replacement.Text = novo_valor
            find_obj.Forward = True
            find_obj.Wrap = 1  # wdFindContinue
            find_obj.Format = False
            find_obj.MatchCase = True  # Case sensitive para os marcadores
            find_obj.MatchWholeWord = False
            find_obj.MatchWildcards = False
            find_obj.MatchSoundsLike = False
            find_obj.MatchAllWordForms = False
            
            # Executar a substituição
            count = 0
            while find_obj.Execute():
                find_obj.Execute(Replace=1)  # Substituir um por vez para contar com precisão
                count += 1
            
            if count > 0:
                total_substituicoes += count
                print(f"Marcador '{marcador}' substituído por '{novo_valor}': {count} ocorrência(s)")
        
        # Aplicar formatação (negrito e alinhamento)
        # Em vez de fazer dentro do loop de substituição, chamar a função separada
        aplicar_formatacao_word(caminho_arquivo)
        
        # Salvar o documento
        doc.Save()
        
        # Fechar o documento
        doc.Close()
        
        # Encerrar o Word
        word.Quit()
        
        if total_substituicoes > 0:
            print(f"Total de substituições realizadas: {total_substituicoes}")
            print("Formatação em negrito e alinhamento aplicados aos campos.")
            return True
        else:
            print("Aviso: Nenhum dos marcadores foi encontrado no documento.")
            return False
    
    except Exception as e:
        print(f"Erro ao modificar o documento com Word: {e}")
        
        # Em caso de erro, garantir que o Word seja fechado
        try:
            if 'doc' in locals() and doc:
                doc.Close(SaveChanges=False)
            if 'word' in locals() and word:
                word.Quit()
        except:
            pass
        
        return False

def modificar_doc_alternativo(caminho_arquivo, substituicoes):
    """
    Método alternativo para modificar um documento .doc
    Usa a conversão temporária para .docx, modificação e conversão de volta
    
    Args:
        caminho_arquivo: Caminho do arquivo a ser modificado
        substituicoes: Dicionário com pares de marcadores e seus novos valores
        
    Returns:
        Boolean indicando se a modificação foi bem-sucedida
    """
    try:
        # Verificar se LibreOffice está disponível
        subprocess.run(["libreoffice", "--version"], stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=False)
        has_libreoffice = True
    except:
        has_libreoffice = False

    try:
        # Criar diretório temporário
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_base = os.path.join(temp_dir, "temp_doc")
            temp_docx = temp_base + ".docx"
            
            print(f"Convertendo o documento para DOCX temporariamente...")
            
            # Usar LibreOffice para converter DOC para DOCX
            if has_libreoffice:
                cmd = [
                    "libreoffice", "--headless", "--convert-to", "docx", 
                    "--outdir", temp_dir, caminho_arquivo
                ]
                process = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                
                if process.returncode != 0:
                    print("Erro ao converter para DOCX usando LibreOffice")
                    return False
            else:
                # Se LibreOffice não estiver disponível, tentar com Word se pywin32 estiver disponível
                if HAS_PYWIN32:
                    try:
                        word = win32com.client.Dispatch("Word.Application")
                        word.Visible = False
                        
                        caminho_completo = os.path.abspath(caminho_arquivo)
                        doc = word.Documents.Open(caminho_completo)
                        
                        doc.SaveAs(os.path.abspath(temp_docx), FileFormat=16)  # 16 = wdFormatDocumentDefault (.docx)
                        doc.Close()
                        word.Quit()
                    except Exception as e:
                        print(f"Erro ao converter para DOCX usando Word: {e}")
                        return False
                else:
                    print("Erro: Não foi possível encontrar LibreOffice ou Word para conversão")
                    return False
            
            # Verificar se o arquivo DOCX temporário foi criado
            if not os.path.exists(temp_docx):
                print("Erro: A conversão para DOCX falhou")
                return False
            
            # Agora temos um arquivo DOCX, podemos modificá-lo com python-docx
            try:
                from docx import Document
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                doc = Document(temp_docx)
                
                # Contador de substituições
                total_substituicoes = 0
                
                # Realizar substituição em parágrafos
                for para in doc.paragraphs:
                    # Verificar se este parágrafo contém algum campo que deve estar em negrito
                    for campo in CAMPOS_NEGRITO:
                        if campo in para.text:
                            # Aplicar negrito a este campo específico
                            runs_original = list(para.runs)
                            
                            # Verificar cada run para encontrar o campo
                            for run in runs_original:
                                if campo in run.text:
                                    # Aplicar negrito ao run que contém o campo
                                    run.bold = True
                    
                    # Verificar se este parágrafo deve estar alinhado à direita
                    for campo in CAMPOS_DIREITA:
                        if campo in para.text:
                            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            break
                    else:
                        # Se não encontrou campo para alinhar à direita, deixar à esquerda
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Agora fazer a substituição de marcadores
                    para_modificado = False
                    texto_original = para.text
                    texto_modificado = texto_original
                    
                    # Verificar e substituir cada marcador no texto
                    for marcador, novo_valor in substituicoes.items():
                        if marcador in texto_modificado:
                            texto_modificado = texto_modificado.replace(marcador, novo_valor)
                            para_modificado = True
                    
                    # Se houve alguma substituição, atualizar o texto mantendo formatação
                    if para_modificado:
                        # Precisamos preservar a formatação original de cada run
                        runs = list(para.runs)
                        texto_atual = ""
                        
                        # Verificar se temos runs para processar
                        if runs:
                            para.clear()
                            
                            # Tentar preservar o máximo de formatação possível
                            for run in runs:
                                run_text = run.text
                                # Se este run contém parte do texto que foi modificado
                                if run_text in texto_original and texto_original.find(run_text) != -1:
                                    start_pos = texto_original.find(run_text, len(texto_atual))
                                    if start_pos != -1:
                                        # Calcular a posição correspondente no texto modificado
                                        relative_pos = len(texto_atual)
                                        run_len = len(run_text)
                                        
                                        # Pegar o texto modificado correspondente a este run
                                        if relative_pos < len(texto_modificado):
                                            end_pos = min(len(texto_modificado), relative_pos + run_len)
                                            new_run_text = texto_modificado[relative_pos:end_pos]
                                            
                                            # Adicionar o run com o texto modificado mantendo formatação
                                            new_run = para.add_run(new_run_text)
                                            new_run.bold = run.bold
                                            new_run.italic = run.italic
                                            new_run.underline = run.underline
                                            if hasattr(run, 'font') and hasattr(run.font, 'color'):
                                                new_run.font.color.rgb = run.font.color.rgb
                                            
                                            texto_atual += run_text
                                        else:
                                            # Se chegamos ao fim do texto modificado, paramos
                                            break
                                else:
                                    # Se não conseguimos processar este run, simplesmente o adicionamos como está
                                    new_run = para.add_run(run_text)
                                    new_run.bold = run.bold
                                    new_run.italic = run.italic
                                    new_run.underline = run.underline
                                    if hasattr(run, 'font') and hasattr(run.font, 'color'):
                                        new_run.font.color.rgb = run.font.color.rgb
                                    
                                    texto_atual += run_text
                        else:
                            # Se não temos runs para processar, simplesmente substituímos o texto
                            para.text = texto_modificado
                        
                        total_substituicoes += 1
                        print(f"Substituição realizada no parágrafo: '{texto_original[:40]}...' -> '{texto_modificado[:40]}...'")
                
                # Realizar substituição em tabelas
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                # Verificar se este parágrafo contém algum campo que deve estar em negrito
                                for campo in CAMPOS_NEGRITO:
                                    if campo in para.text:
                                        # Aplicar negrito a este campo específico
                                        runs_original = list(para.runs)
                                        
                                        # Verificar cada run para encontrar o campo
                                        for run in runs_original:
                                            if campo in run.text:
                                                # Aplicar negrito ao run que contém o campo
                                                run.bold = True
                                
                                # Verificar se este parágrafo deve estar alinhado à direita
                                for campo in CAMPOS_DIREITA:
                                    if campo in para.text:
                                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                        break
                                else:
                                    # Se não encontrou campo para alinhar à direita, deixar à esquerda
                                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                
                                # Agora fazer a substituição de marcadores
                                para_modificado = False
                                texto_original = para.text
                                texto_modificado = texto_original
                                
                                # Verificar e substituir cada marcador no texto
                                for marcador, novo_valor in substituicoes.items():
                                    if marcador in texto_modificado:
                                        texto_modificado = texto_modificado.replace(marcador, novo_valor)
                                        para_modificado = True
                                
                                # Se houve alguma substituição, atualizar o texto mantendo formatação
                                if para_modificado:
                                    # Similarmente ao processo para parágrafos normais
                                    runs = list(para.runs)
                                    if runs:
                                        para.clear()
                                        for run in runs:
                                            new_run = para.add_run(run.text.replace(marcador, novo_valor) if marcador in run.text else run.text)
                                            new_run.bold = run.bold
                                            new_run.italic = run.italic
                                            new_run.underline = run.underline
                                            if hasattr(run, 'font') and hasattr(run.font, 'color'):
                                                new_run.font.color.rgb = run.font.color.rgb
                                    else:
                                        para.text = texto_modificado
                                    
                                    total_substituicoes += 1
                                    print(f"Substituição realizada na tabela: '{texto_original[:40]}...' -> '{texto_modificado[:40]}...'")
                
                if total_substituicoes == 0:
                    print("Aviso: Nenhum dos marcadores foi encontrado no documento.")
                    return False
                
                # Salvar o arquivo DOCX modificado
                doc.save(temp_docx)
                
                # Converter de volta para DOC
                print("Convertendo de volta para o formato DOC...")
                
                # Substituir o arquivo original pelo modificado
                temp_doc = temp_base + ".doc"
                
                if has_libreoffice:
                    cmd = [
                        "libreoffice", "--headless", "--convert-to", "doc", 
                        "--outdir", temp_dir, temp_docx
                    ]
                    process = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                    
                    if process.returncode != 0:
                        print("Erro ao converter de volta para DOC")
                        return False
                    
                    # Copiar o arquivo temporário para o destino final
                    shutil.copy2(temp_doc, caminho_arquivo)
                else:
                    # Se LibreOffice não estiver disponível, tentar com Word
                    if HAS_PYWIN32:
                        try:
                            word = win32com.client.Dispatch("Word.Application")
                            word.Visible = False
                            
                            doc = word.Documents.Open(os.path.abspath(temp_docx))
                            
                            doc.SaveAs(os.path.abspath(caminho_arquivo), FileFormat=0)  # 0 = wdFormatDocument (.doc)
                            doc.Close()
                            word.Quit()
                        except Exception as e:
                            print(f"Erro ao converter de volta para DOC usando Word: {e}")
                            return False
                    else:
                        print("Erro: Não foi possível encontrar LibreOffice ou Word para conversão de volta")
                        return False
                
                # Se estamos usando Word, aplicar formatação adicional diretamente com Word
                if HAS_PYWIN32 and not has_libreoffice:
                    aplicar_formatacao_word(caminho_arquivo)
                
                print(f"Documento modificado e formatado com sucesso: {caminho_arquivo}")
                return True
                
            except ImportError:
                print("Erro: A biblioteca 'python-docx' é necessária para esta função.")
                print("Para instalar, execute: pip install python-docx")
                return False
            except Exception as e:
                print(f"Erro ao modificar o arquivo DOCX: {e}")
                return False
    
    except Exception as e:
        print(f"Erro no processo de conversão/modificação: {e}")
        return False

def perguntar_metodo():
    """
    Pergunta ao usuário qual método usar para modificar o documento
    
    Returns:
        String com o método escolhido
    """
    print("\nEscolha o método para modificar o documento:")
    print("1. Usar Microsoft Word (recomendado para máxima compatibilidade)")
    print("2. Usar método alternativo (conversão DOCX temporária)")
    
    escolha = input("\nDigite o número da opção desejada (1-2): ")
    
    if escolha == "1":
        if not HAS_PYWIN32:
            print("A biblioteca 'pywin32' não está instalada. Usando método alternativo.")
            return "alternativo"
        return "word"
    else:
        return "alternativo"

def obter_valores_para_marcadores():
    """
    Solicita ao usuário os valores para cada marcador
    
    Returns:
        Dicionário com os pares marcador-valor
    """
    substituicoes = {}
    data_atual = datetime.now().strftime("%d/%m/%Y")
    
    print("\n--- Substituição de Marcadores ---")
    print("Digite os valores para substituir cada marcador (ou pressione Enter para usar o valor sugerido):")
    
    # Data da Solicitação (sugestão: data atual)
    valor = input(f"\n[Data da Solicitação] (sugestão: {data_atual}): ")
    substituicoes["[Data da Solicitação]"] = valor if valor else data_atual
    
    # Data de Execução (sugestão: data atual)
    valor = input(f"\n[Data de Execução] (sugestão: {data_atual}): ")
    substituicoes["[Data de Execução]"] = valor if valor else data_atual
    
    # Localização
    valor = input("\n[Localização]: ")
    if valor:
        substituicoes["[Localização]"] = valor
    
    # O.S
    valor = input("\n[O.S]: ")
    if valor:
        substituicoes["[O.S]"] = valor
    
    # SERVIÇO A REALIZAR
    valor = input("\n[SERVIÇO A REALIZAR]: ")
    if valor:
        substituicoes["[SERVIÇO A REALIZAR]"] = valor
    
    # OBSERVAÇÃO
    valor = input("\n[OBSERVAÇÃO]: ")
    if valor:
        substituicoes["[OBSERVAÇÃO]"] = valor
    
    # Mostrar resumo das substituições
    print("\nSubstituições que serão realizadas:")
    for marcador, valor in substituicoes.items():
        print(f"{marcador} -> {valor}")
    
    # Confirmar se deseja continuar
    confirmacao = input("\nDeseja continuar com essas substituições? (S/N): ")
    if confirmacao.upper() != "S":
        print("Operação cancelada pelo usuário.")
        sys.exit(0)
    
    return substituicoes

if __name__ == "__main__":
    # Configuração dos caminhos
    caminho_origem = "report.doc"
    
    # Verificar se o arquivo de origem existe
    if not os.path.exists(caminho_origem):
        print(f"Erro: O arquivo {caminho_origem} não foi encontrado.")
        # Solicitar o caminho do arquivo
        caminho_origem = input("Por favor, digite o caminho completo do arquivo .doc que deseja modificar: ")
        
        # Verificar novamente
        if not os.path.exists(caminho_origem):
            print(f"Erro: O arquivo {caminho_origem} não foi encontrado.")
            sys.exit(1)
    
    # Gerar nome para arquivo de trabalho baseado no nome original
    nome_base = os.path.splitext(os.path.basename(caminho_origem))[0]
    caminho_trabalho = f"{nome_base}_trabalho_temp.doc"
    
    # Fase 1: Criar uma cópia exata de trabalho
    print(f"Processando o arquivo {caminho_origem}...")
    print(f"Tamanho do arquivo original: {os.path.getsize(caminho_origem)} bytes")
    
    # Replicar o documento
    sucesso = replicar_documento(caminho_origem, caminho_trabalho)
    
    if not sucesso or not os.path.exists(caminho_trabalho):
        print(f"Erro: Não foi possível criar a cópia de trabalho.")
        sys.exit(1)
    
    # Verificar se o conteúdo é idêntico
    if not verificar_conteudo(caminho_origem, caminho_trabalho):
        print("Aviso: A cópia de trabalho não é idêntica ao original. Isso pode causar problemas.")
        continuar = input("Deseja continuar mesmo assim? (S/N): ")
        if continuar.upper() != "S":
            print("Operação cancelada pelo usuário.")
            sys.exit(0)
    
    # Fase 2: Solicitar valores para os marcadores
    substituicoes = obter_valores_para_marcadores()
    
    # Fase 3: Escolher método de modificação
    metodo = perguntar_metodo()
    
    # Fase 4: Aplicar a modificação
    if metodo == "word":
        print("\nModificando o documento usando Microsoft Word...")
        sucesso = modificar_doc_com_word(caminho_trabalho, substituicoes)
    else:
        print("\nModificando o documento usando método alternativo...")
        sucesso = modificar_doc_alternativo(caminho_trabalho, substituicoes)
    
    # Fase 5: Verificar resultado e finalizar
    if sucesso:
        print("\nVerificando o resultado...")
        if verificar_conteudo(caminho_origem, caminho_trabalho):
            print("Aviso: O documento modificado é idêntico ao original. As substituições podem não ter sido aplicadas.")
        else:
            print("\nModificação concluída com sucesso!")
            print(f"O documento modificado está em: {caminho_trabalho}")
            
            # Perguntar se deseja substituir o original
            backup = f"{nome_base}_original_backup.doc"
            substituir = input("\nDeseja substituir o arquivo original? (S/N): ")
            if substituir.upper() == "S":
                # Fazer backup do original
                shutil.copy2(caminho_origem, backup)
                print(f"Backup do arquivo original criado em: {backup}")
                
                # Substituir o original
                shutil.copy2(caminho_trabalho, caminho_origem)
                print(f"Arquivo original substituído.")
    else:
        print("\nA modificação falhou. A cópia de trabalho pode estar corrompida.") 