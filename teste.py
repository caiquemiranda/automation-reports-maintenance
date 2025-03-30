import os
import shutil
import re
import tempfile
import subprocess
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def copiar_formatacao_paragrafo(paragrafo_origem, paragrafo_destino):
    """
    Copia a formatação de um parágrafo para outro
    
    Args:
        paragrafo_origem: Parágrafo de origem
        paragrafo_destino: Parágrafo de destino
    """
    # Copiar alinhamento
    paragrafo_destino.alignment = paragrafo_origem.alignment
    
    # Copiar recuos e espaçamentos
    paragrafo_destino.paragraph_format.left_indent = paragrafo_origem.paragraph_format.left_indent
    paragrafo_destino.paragraph_format.right_indent = paragrafo_origem.paragraph_format.right_indent
    paragrafo_destino.paragraph_format.first_line_indent = paragrafo_origem.paragraph_format.first_line_indent
    paragrafo_destino.paragraph_format.space_before = paragrafo_origem.paragraph_format.space_before
    paragrafo_destino.paragraph_format.space_after = paragrafo_origem.paragraph_format.space_after
    paragrafo_destino.paragraph_format.line_spacing = paragrafo_origem.paragraph_format.line_spacing
    
    # Copiar estilo de lista, se houver
    if hasattr(paragrafo_origem, 'style') and paragrafo_origem.style:
        try:
            paragrafo_destino.style = paragrafo_origem.style
        except Exception:
            print("Aviso: Não foi possível copiar o estilo do parágrafo")

def copiar_formatacao_run(run_origem, run_destino):
    """
    Copia a formatação de um run para outro
    
    Args:
        run_origem: Run de origem (trecho de texto com formatação uniforme)
        run_destino: Run de destino
    """
    # Copiar fonte e tamanho
    run_destino.font.name = run_origem.font.name
    run_destino.font.size = run_origem.font.size
    
    # Copiar estilo (negrito, itálico, etc)
    run_destino.font.bold = run_origem.font.bold
    run_destino.font.italic = run_origem.font.italic
    run_destino.font.underline = run_origem.font.underline
    run_destino.font.strike = run_origem.font.strike  # Texto tachado
    run_destino.font.subscript = run_origem.font.subscript  # Subscrito
    run_destino.font.superscript = run_origem.font.superscript  # Sobrescrito
    
    # Copiar cor
    if run_origem.font.color.rgb:
        run_destino.font.color.rgb = run_origem.font.color.rgb
    
    # Copiar realce
    run_destino.font.highlight_color = run_origem.font.highlight_color
    
    # Copiar espaçamento entre caracteres
    if hasattr(run_origem.font, 'character_spacing'):
        run_destino.font.character_spacing = run_origem.font.character_spacing

def copiar_tabela(tabela_origem, documento_destino):
    """
    Copia uma tabela para o documento de destino
    
    Args:
        tabela_origem: Tabela de origem
        documento_destino: Documento de destino
        
    Returns:
        Tabela copiada
    """
    # Criar nova tabela com o mesmo número de linhas e colunas
    linhas = len(tabela_origem.rows)
    colunas = len(tabela_origem.columns)
    
    tabela_destino = documento_destino.add_table(rows=linhas, cols=colunas)
    
    # Copiar estilo da tabela
    if tabela_origem.style:
        try:
            tabela_destino.style = tabela_origem.style
        except Exception:
            print("Aviso: Não foi possível copiar o estilo da tabela")
    
    # Copiar largura das colunas (se possível)
    try:
        for i in range(colunas):
            for cell in tabela_destino.column_cells(i):
                cell.width = tabela_origem.column_cells(i)[0].width
    except Exception:
        print("Aviso: Não foi possível copiar a largura exata das colunas")
    
    # Copiar conteúdo e formatação célula por célula
    for i, linha in enumerate(tabela_origem.rows):
        # Definir altura da linha (se possível)
        try:
            tabela_destino.rows[i].height = linha.height
        except:
            pass
            
        for j, celula in enumerate(linha.cells):
            celula_destino = tabela_destino.cell(i, j)
            
            # Copiar mesclagem (se aplicável)
            if celula.merge_origin:
                try:
                    # Identificar células a mesclar na tabela de destino
                    # Isso é uma simplificação, o ideal seria identificar o intervalo exato
                    pass
                except:
                    print("Aviso: Não foi possível copiar a mesclagem da célula")
            
            # Copiar o conteúdo
            # Remover parágrafo padrão se não for o primeiro
            if len(celula.paragraphs) > 1:
                celula_destino.paragraphs.clear()
                
            for paragrafo_origem in celula.paragraphs:
                if len(celula_destino.paragraphs) > 0 and celula_destino.paragraphs[0].text == "":
                    paragrafo_destino = celula_destino.paragraphs[0]
                else:
                    paragrafo_destino = celula_destino.add_paragraph()
                
                # Copiar o texto e formatação
                for run_origem in paragrafo_origem.runs:
                    run_destino = paragrafo_destino.add_run(run_origem.text)
                    copiar_formatacao_run(run_origem, run_destino)
                
                copiar_formatacao_paragrafo(paragrafo_origem, paragrafo_destino)
            
            # Copiar bordas (simplificado)
            try:
                if hasattr(celula, 'border_top') and celula.border_top:
                    celula_destino.border_top = celula.border_top
                if hasattr(celula, 'border_bottom') and celula.border_bottom:
                    celula_destino.border_bottom = celula.border_bottom
                if hasattr(celula, 'border_left') and celula.border_left:
                    celula_destino.border_left = celula.border_left
                if hasattr(celula, 'border_right') and celula.border_right:
                    celula_destino.border_right = celula.border_right
            except:
                print("Aviso: Não foi possível copiar as bordas da célula")
    
    return tabela_destino

def copiar_imagem(imagem_origem, documento_destino):
    """
    Copia uma imagem para o documento de destino
    
    Args:
        imagem_origem: Imagem de origem (bloco de imagem)
        documento_destino: Documento de destino
    """
    try:
        # Em implementações mais complexas, seria necessário extrair a imagem e reinseri-la
        # Esta é uma implementação simplificada
        pass
    except Exception as e:
        print(f"Erro ao copiar imagem: {e}")

def modificar_servico(doc, servico_antigo, servico_novo):
    """
    Modifica o texto do serviço no documento
    
    Args:
        doc: Documento a ser modificado
        servico_antigo: Texto do serviço a ser substituído ou padrão regex
        servico_novo: Novo texto do serviço
        
    Returns:
        Boolean indicando se a substituição foi realizada
    """
    substituicao_realizada = False
    
    # Procurar em todos os parágrafos
    for paragrafo in doc.paragraphs:
        texto_paragrafo = paragrafo.text
        
        # Verificar se o texto do serviço está no parágrafo (exato ou por regex)
        if isinstance(servico_antigo, str) and servico_antigo in texto_paragrafo:
            # Substituição de texto exato
            novo_texto = texto_paragrafo.replace(servico_antigo, servico_novo)
            
            # Limpar o parágrafo e adicionar o novo texto com a mesma formatação
            runs_formatacao = [run for run in paragrafo.runs]
            paragrafo.clear()
            
            # Se tivermos múltiplos runs, tentamos preservar a formatação correspondente
            if runs_formatacao:
                # Encontrar onde o texto original estava
                inicio_servico = texto_paragrafo.find(servico_antigo)
                fim_servico = inicio_servico + len(servico_antigo)
                
                # Reconstruir o parágrafo com o texto modificado
                pos_atual = 0
                for run in runs_formatacao:
                    run_len = len(run.text)
                    # Se este run está completamente antes do serviço
                    if pos_atual + run_len <= inicio_servico:
                        novo_run = paragrafo.add_run(run.text)
                        copiar_formatacao_run(run, novo_run)
                    # Se este run contém o início do serviço
                    elif pos_atual <= inicio_servico < pos_atual + run_len:
                        # Texto antes do serviço
                        if inicio_servico > pos_atual:
                            texto_antes = run.text[:inicio_servico-pos_atual]
                            novo_run = paragrafo.add_run(texto_antes)
                            copiar_formatacao_run(run, novo_run)
                        
                        # Adicionar o novo serviço com a formatação deste run
                        novo_run = paragrafo.add_run(servico_novo)
                        copiar_formatacao_run(run, novo_run)
                        
                        # Se o serviço termina neste run, adicionar o resto
                        if fim_servico < pos_atual + run_len:
                            texto_depois = run.text[fim_servico-pos_atual:]
                            novo_run = paragrafo.add_run(texto_depois)
                            copiar_formatacao_run(run, novo_run)
                    # Se este run está completamente após o serviço
                    elif pos_atual >= fim_servico:
                        novo_run = paragrafo.add_run(run.text)
                        copiar_formatacao_run(run, novo_run)
                    # Caso contrário, este run está dentro ou atravessa o serviço, então pular
                    
                    pos_atual += run_len
            else:
                # Se não conseguirmos preservar a formatação, apenas substituir o texto
                novo_run = paragrafo.add_run(novo_texto)
            
            substituicao_realizada = True
            print(f"Substituição realizada no parágrafo: '{texto_paragrafo}' -> '{novo_texto}'")
            
        elif isinstance(servico_antigo, re.Pattern) and servico_antigo.search(texto_paragrafo):
            # Substituição por regex
            match = servico_antigo.search(texto_paragrafo)
            if match:
                novo_texto = servico_antigo.sub(servico_novo, texto_paragrafo)
                
                # Limpar o parágrafo e adicionar o novo texto
                runs_formatacao = [run for run in paragrafo.runs]
                paragrafo.clear()
                
                # Tenta preservar formatação (simplificado para regex)
                if runs_formatacao:
                    # Usar a formatação do primeiro run para todo o texto
                    novo_run = paragrafo.add_run(novo_texto)
                    copiar_formatacao_run(runs_formatacao[0], novo_run)
                else:
                    novo_run = paragrafo.add_run(novo_texto)
                
                substituicao_realizada = True
                print(f"Substituição realizada por regex: '{texto_paragrafo}' -> '{novo_texto}'")
    
    # Procurar em todas as tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    texto_paragrafo = paragrafo.text
                    
                    # Verificar se o texto do serviço está no parágrafo (exato ou por regex)
                    if isinstance(servico_antigo, str) and servico_antigo in texto_paragrafo:
                        # Substituição de texto exato em tabela
                        novo_texto = texto_paragrafo.replace(servico_antigo, servico_novo)
                        
                        # Preservar formatação (simplificado para células de tabela)
                        runs_formatacao = [run for run in paragrafo.runs]
                        paragrafo.clear()
                        
                        if runs_formatacao:
                            # Usar a formatação do primeiro run
                            novo_run = paragrafo.add_run(novo_texto)
                            copiar_formatacao_run(runs_formatacao[0], novo_run)
                        else:
                            novo_run = paragrafo.add_run(novo_texto)
                        
                        substituicao_realizada = True
                        print(f"Substituição realizada na tabela: '{texto_paragrafo}' -> '{novo_texto}'")
                        
                    elif isinstance(servico_antigo, re.Pattern) and servico_antigo.search(texto_paragrafo):
                        # Substituição por regex em tabela
                        match = servico_antigo.search(texto_paragrafo)
                        if match:
                            novo_texto = servico_antigo.sub(servico_novo, texto_paragrafo)
                            
                            # Simplificado para células de tabela
                            runs_formatacao = [run for run in paragrafo.runs]
                            paragrafo.clear()
                            
                            if runs_formatacao:
                                # Usar a formatação do primeiro run para todo o texto
                                novo_run = paragrafo.add_run(novo_texto)
                                copiar_formatacao_run(runs_formatacao[0], novo_run)
                            else:
                                novo_run = paragrafo.add_run(novo_texto)
                            
                            substituicao_realizada = True
                            print(f"Substituição realizada por regex na tabela: '{texto_paragrafo}' -> '{novo_texto}'")
    
    return substituicao_realizada

def substituir_texto_binario(caminho_origem, caminho_destino, substituicoes):
    """
    Substitui texto em arquivos binários (.doc) usando um método alternativo
    
    Args:
        caminho_origem: Caminho do arquivo de origem
        caminho_destino: Caminho do arquivo de destino
        substituicoes: Dicionário com pares de strings a serem substituídas
        
    Returns:
        Boolean indicando se a substituição foi realizada
    """
    try:
        # Primeiro, fazer uma cópia exata
        shutil.copy2(caminho_origem, caminho_destino)
        
        # Ler o arquivo como binário
        with open(caminho_destino, 'rb') as f:
            conteudo = f.read()
        
        # Realizar substituições binárias
        modificacoes_realizadas = False
        novo_conteudo = conteudo
        
        for texto_antigo, texto_novo in substituicoes.items():
            # Ignorar padrões regex, eles não funcionam bem em binário
            if isinstance(texto_antigo, str):
                # Converter strings para bytes com diversas codificações comuns
                encodings = ['utf-8', 'latin1', 'utf-16', 'utf-16le', 'utf-16be']
                
                for encoding in encodings:
                    try:
                        texto_antigo_bytes = texto_antigo.encode(encoding)
                        texto_novo_bytes = texto_novo.encode(encoding)
                        
                        # Garantir que o texto novo tenha o mesmo tamanho ou seja menor
                        if len(texto_novo_bytes) > len(texto_antigo_bytes):
                            # Preencher com espaços para manter o mesmo tamanho
                            texto_novo_bytes = texto_novo.ljust(len(texto_antigo)).encode(encoding)
                        
                        # Verificar se o texto antigo existe no conteúdo
                        if texto_antigo_bytes in novo_conteudo:
                            # Substituir todas as ocorrências
                            novo_conteudo = novo_conteudo.replace(texto_antigo_bytes, texto_novo_bytes)
                            modificacoes_realizadas = True
                            print(f"Substituição binária realizada com encoding {encoding}: '{texto_antigo}' -> '{texto_novo}'")
                    except Exception as e:
                        print(f"Erro ao tentar substituição binária com encoding {encoding}: {e}")
        
        # Salvar se houve modificações
        if modificacoes_realizadas:
            with open(caminho_destino, 'wb') as f:
                f.write(novo_conteudo)
            print(f"Arquivo modificado e salvo em: {caminho_destino}")
            return True
        else:
            print("Aviso: Nenhuma substituição binária foi realizada.")
            return False
            
    except Exception as e:
        print(f"Erro ao tentar substituição binária: {e}")
        return False

def buscar_extrair_texto(caminho_arquivo):
    """
    Tenta extrair texto de um arquivo DOC/DOCX usando métodos alternativos
    
    Args:
        caminho_arquivo: Caminho do arquivo
        
    Returns:
        String com o texto extraído ou None
    """
    texto_extraido = ""
    
    try:
        # Tentar carregar como DOCX primeiro
        try:
            doc = Document(caminho_arquivo)
            
            # Extrair texto de parágrafos
            for paragrafo in doc.paragraphs:
                texto_extraido += paragrafo.text + "\n"
            
            # Extrair texto de tabelas
            for tabela in doc.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        for paragrafo in celula.paragraphs:
                            texto_extraido += paragrafo.text + " "
                    texto_extraido += "\n"
            
            if texto_extraido:
                return texto_extraido
            
        except Exception as e:
            print(f"Aviso: Não foi possível extrair texto com python-docx: {e}")
        
        # Se falhar, tentar abordagem alternativa para arquivos binários
        if caminho_arquivo.lower().endswith('.doc'):
            print("Tentando extrair texto do arquivo binário .doc usando método alternativo...")
            
            # Ler como binário e procurar por strings
            with open(caminho_arquivo, 'rb') as f:
                conteudo = f.read()
            
            # Tentar extrair strings legíveis com várias codificações
            for encoding in ['utf-8', 'latin1', 'utf-16le']:
                try:
                    # Decodificar e manter apenas caracteres legíveis
                    texto_decodificado = conteudo.decode(encoding, errors='ignore')
                    texto_limpo = ''.join(c if c.isprintable() else ' ' for c in texto_decodificado)
                    
                    # Remover múltiplos espaços em branco
                    texto_limpo = re.sub(r'\s+', ' ', texto_limpo).strip()
                    
                    # Se encontramos texto substancial, retornar
                    if len(texto_limpo) > 100:  # Arbitrário, ajustar conforme necessário
                        return texto_limpo
                except Exception:
                    continue
    
    except Exception as e:
        print(f"Erro ao tentar extrair texto: {e}")
    
    return None

def replicar_documento(caminho_origem, caminho_destino):
    """
    Replica o documento de origem para o de destino usando apenas cópia binária
    
    Args:
        caminho_origem: Caminho do arquivo de origem
        caminho_destino: Caminho do arquivo de destino
        
    Returns:
        Boolean indicando se a operação foi bem-sucedida
    """
    try:
        # Realizar cópia direta (binária) para evitar problemas de corrupção
        print(f"Realizando cópia binária do arquivo {caminho_origem}...")
        shutil.copy2(caminho_origem, caminho_destino)
        print(f"Arquivo copiado com sucesso para: {caminho_destino}")
        return True
    except Exception as e:
        print(f"Erro ao replicar o documento: {e}")
        return False

def verificar_conteudo(caminho_origem, caminho_destino):
    """
    Verifica se o conteúdo dos arquivos é idêntico
    
    Args:
        caminho_origem: Caminho do arquivo de origem
        caminho_destino: Caminho do arquivo de destino
        
    Returns:
        Boolean indicando se os arquivos são idênticos
    """
    try:
        with open(caminho_origem, 'rb') as f1, open(caminho_destino, 'rb') as f2:
            conteudo1 = f1.read()
            conteudo2 = f2.read()
            
            if conteudo1 == conteudo2:
                print("Verificação: Os arquivos são IDÊNTICOS!")
                return True
            else:
                print("Verificação: Os arquivos são DIFERENTES!")
                tamanho1 = len(conteudo1)
                tamanho2 = len(conteudo2)
                print(f"Tamanho do arquivo original: {tamanho1} bytes")
                print(f"Tamanho do arquivo replicado: {tamanho2} bytes")
                
                if tamanho1 == tamanho2:
                    print("Os arquivos têm o mesmo tamanho, mas conteúdo diferente")
                else:
                    print(f"Diferença de tamanho: {abs(tamanho1 - tamanho2)} bytes")
                
                return False
    except Exception as e:
        print(f"Erro ao verificar conteúdo: {e}")
        return False

def ler_servico_atual(caminho_arquivo):
    """
    Tenta ler o serviço atual do documento
    
    Args:
        caminho_arquivo: Caminho do arquivo
        
    Returns:
        String com o serviço encontrado ou None
    """
    try:
        # Extrair todo o texto do documento
        texto_extraido = buscar_extrair_texto(caminho_arquivo)
        
        if texto_extraido:
            # Padrão para encontrar o serviço
            padrao_servico = re.compile(r'SERVIÇO:.*?[\(\)]', re.IGNORECASE)
            
            # Procurar no texto extraído
            match = padrao_servico.search(texto_extraido)
            if match:
                return match.group(0)
            
            # Tentar outro padrão mais flexível
            padrao_alternativo = re.compile(r'SERVIÇO:.*?[0-9]', re.IGNORECASE)
            match = padrao_alternativo.search(texto_extraido)
            if match:
                return match.group(0)
        
        print("Aviso: Não foi encontrado texto de serviço no documento")
        return None
        
    except Exception as e:
        print(f"Erro ao ler serviço atual: {e}")
        return None

if __name__ == "__main__":
    # Configuração dos caminhos
    caminho_origem = "report.doc"
    caminho_destino = "report_replicado.doc"
    
    # Verificar se o arquivo de origem existe
    if not os.path.exists(caminho_origem):
        print(f"Erro: O arquivo {caminho_origem} não foi encontrado.")
    else:
        print(f"Processando o arquivo {caminho_origem}...")
        print(f"Tamanho do arquivo original: {os.path.getsize(caminho_origem)} bytes")
        
        # Replicar o documento
        sucesso = replicar_documento(caminho_origem, caminho_destino)
        
        # Verificar se o arquivo foi criado
        if os.path.exists(caminho_destino):
            print(f"Tamanho do arquivo replicado: {os.path.getsize(caminho_destino)} bytes")
            
            # Verificar se o conteúdo é idêntico
            verificar_conteudo(caminho_origem, caminho_destino)
        else:
            print(f"Erro: O arquivo {caminho_destino} não foi criado corretamente.")
