from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
from docx import Document
import tempfile
from models import MaintenanceReport

app = Flask(__name__)
CORS(app)

@app.route('/api/health', methods=['GET'])
def health_check():
    return jsonify({"status": "online", "service": "maintenance-reports-api"})

@app.route('/api/generate-report', methods=['POST'])
def generate_report():
    try:
        data = request.json
        report = MaintenanceReport(**data)
        
        # Criar documento Word
        doc = Document()
        doc.add_heading('Relatório de Manutenção', 0)
        
        # Adicionar informações do relatório
        doc.add_heading('Informações Gerais', level=1)
        doc.add_paragraph(f'Equipamento: {report.equipment}')
        doc.add_paragraph(f'Data: {report.date}')
        doc.add_paragraph(f'Técnico: {report.technician}')
        
        doc.add_heading('Descrição do Serviço', level=1)
        doc.add_paragraph(report.service_description)
        
        doc.add_heading('Peças Utilizadas', level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Cabeçalho da tabela
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Nome'
        header_cells[1].text = 'Quantidade'
        header_cells[2].text = 'Observações'
        
        # Adicionar peças na tabela
        for part in report.parts_used:
            row_cells = table.add_row().cells
            row_cells[0].text = part.name
            row_cells[1].text = str(part.quantity)
            row_cells[2].text = part.notes
        
        # Salvar o documento em um arquivo temporário
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_filename = temp_file.name
        doc.save(temp_filename)
        temp_file.close()
        
        # Enviar o arquivo para o cliente
        return send_file(
            temp_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'relatorio_{report.equipment}_{report.date}.docx'
        )
        
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    
    finally:
        # Limpar o arquivo temporário
        if 'temp_filename' in locals():
            try:
                os.unlink(temp_filename)
            except:
                pass

if __name__ == '__main__':
    app.run(debug=True, port=5000) 