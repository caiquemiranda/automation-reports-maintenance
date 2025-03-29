from docx import Document
import tempfile
import os
from typing import Dict, Any, Optional, List
from ..models import MaintenanceReport, Part

class DocumentService:
    """Serviço para geração de documentos Word."""
    
    @staticmethod
    def generate_report_document(report: MaintenanceReport) -> str:
        """
        Gera um documento Word a partir de um relatório.
        
        Args:
            report: O relatório a ser convertido em documento
            
        Returns:
            Caminho do arquivo temporário com o documento gerado
            
        Raises:
            Exception: Se ocorrer um erro durante a geração do documento
        """
        try:
            # Criar documento Word
            doc = Document()
            doc.add_heading('Relatório de Manutenção', 0)
            
            # Adicionar informações do relatório
            doc.add_heading('Informações Gerais', level=1)
            doc.add_paragraph(f'Equipamento: {report.equipment}')
            doc.add_paragraph(f'Data: {report.date}')
            doc.add_paragraph(f'Técnico: {report.technician}')
            
            doc.add_heading('Descrição do Serviço', level=1)
            doc.add_paragraph(report.service_description)
            
            doc.add_heading('Peças Utilizadas', level=1)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            # Cabeçalho da tabela
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Nome'
            header_cells[1].text = 'Quantidade'
            header_cells[2].text = 'Observações'
            
            # Adicionar peças na tabela
            for part in report.parts_used:
                row_cells = table.add_row().cells
                row_cells[0].text = part.name
                row_cells[1].text = str(part.quantity)
                row_cells[2].text = part.notes or ""
            
            # Adicionar observações, se houver
            if report.observations:
                doc.add_heading('Observações Adicionais', level=1)
                doc.add_paragraph(report.observations)
            
            # Salvar o documento em um arquivo temporário
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            temp_filename = temp_file.name
            doc.save(temp_filename)
            temp_file.close()
            
            return temp_filename
            
        except Exception as e:
            # Se ocorrer erro, limpar arquivo temporário se ele existir
            if 'temp_filename' in locals():
                try:
                    os.unlink(temp_filename)
                except:
                    pass
            raise Exception(f"Erro ao gerar documento: {str(e)}")
    
    @staticmethod
    def cleanup_temp_file(file_path: str) -> None:
        """
        Remove um arquivo temporário.
        
        Args:
            file_path: Caminho do arquivo a ser removido
        """
        try:
            if os.path.exists(file_path):
                os.unlink(file_path)
        except Exception:
            # Ignore errors when cleaning up temporary files
            pass 